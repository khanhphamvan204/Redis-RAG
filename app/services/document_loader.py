# # app/services/document_loader.py
# import os
# import logging
# import cv2
# import numpy as np
# import pytesseract
# from pdf2image import convert_from_path
# from PIL import Image
# from langchain.docstore.document import Document
# from langchain_community.document_loaders import (
#     TextLoader,
#     Docx2txtLoader,
#     CSVLoader,
#     UnstructuredExcelLoader
# )

# logger = logging.getLogger(__name__)

# # Cấu hình đường dẫn Tesseract
# base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
# pytesseract.pytesseract.tesseract_cmd = os.path.join(base_dir, "app", "models", "Tesseract-OCR", "tesseract.exe")
# tessdata_dir = os.path.join(base_dir, "app", "models", "Tesseract-OCR", "tessdata")

# # Global variables for OCR
# ocr_initialized = False

# def init_tesseract():
#     """Khởi tạo và kiểm tra Tesseract OCR"""
#     global ocr_initialized
    
#     if ocr_initialized:
#         return True
    
#     try:
#         logger.info("Initializing Tesseract OCR...")
        
#         # Kiểm tra xem tesseract executable có tồn tại không
#         if not os.path.exists(pytesseract.pytesseract.tesseract_cmd):
#             logger.error(f"Tesseract executable not found at: {pytesseract.pytesseract.tesseract_cmd}")
#             return False
        
#         # Kiểm tra tessdata directory
#         if not os.path.exists(tessdata_dir):
#             logger.error(f"Tessdata directory not found at: {tessdata_dir}")
#             return False
        
#         # Kiểm tra các file ngôn ngữ có sẵn
#         available_langs = []
#         for lang_file in ['vie.traineddata', 'eng.traineddata', 'osd.traineddata']:
#             lang_path = os.path.join(tessdata_dir, lang_file)
#             if os.path.exists(lang_path):
#                 available_langs.append(lang_file.replace('.traineddata', ''))
        
#         if not available_langs:
#             logger.error(f"No language data files found in: {tessdata_dir}")
#             return False
        
#         logger.info(f"Available languages: {available_langs}")
        
#         # Set environment variable cho tessdata
#         tessdata_path = tessdata_dir.replace('\\', '/')
#         os.environ['TESSDATA_PREFIX'] = tessdata_path
        
#         # Test với tiếng Việt hoặc tiếng Anh
#         test_lang = 'vie' if 'vie' in available_langs else 'eng'
#         if test_lang not in available_langs:
#             logger.error("Neither Vietnamese nor English language files found")
#             return False
        
#         # Test OCR
#         test_img = Image.new('RGB', (200, 50), color='white')
#         from PIL import ImageDraw
#         draw = ImageDraw.Draw(test_img)
#         draw.text((10, 10), "TEST", fill='black')
        
#         config = f'--tessdata-dir {tessdata_path} -l {test_lang}'
#         test_result = pytesseract.image_to_string(test_img, config=config)
        
#         ocr_initialized = True
#         logger.info(f"Tesseract OCR initialized successfully with {test_lang}")
#         return True
        
#     except Exception as e:
#         logger.error(f"Failed to initialize Tesseract OCR: {str(e)}")
#         return False

# def preprocess_image_for_ocr(image):
#     """
#     Tiền xử lý hình ảnh để cải thiện OCR
#     """
#     try:
#         # Convert PIL Image to numpy array
#         img_array = np.array(image)
        
#         # Convert RGB to BGR for OpenCV
#         if len(img_array.shape) == 3:
#             img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
#             gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
#         else:
#             gray = img_array
        
#         # Apply CLAHE (Contrast Limited Adaptive Histogram Equalization)
#         clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
#         enhanced = clahe.apply(gray)
        
#         # Denoise
#         denoised = cv2.medianBlur(enhanced, 3)
        
#         # Threshold
#         _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
#         return Image.fromarray(thresh)
        
#     except Exception as e:
#         logger.warning(f"Image preprocessing failed: {str(e)}, using original image")
#         return image

# def extract_text_with_tesseract(image, page_num=1):
#     """Trích xuất văn bản từ hình ảnh bằng Tesseract OCR"""
#     if not init_tesseract():
#         logger.error("Cannot initialize Tesseract OCR")
#         return ""
    
#     try:
#         # Preprocess image
#         processed_img = preprocess_image_for_ocr(image)
        
#         # Prefer Vietnamese, fallback to English
#         selected_lang = 'vie'
#         if not os.path.exists(os.path.join(tessdata_dir, f"{selected_lang}.traineddata")):
#             selected_lang = 'eng'
#             if not os.path.exists(os.path.join(tessdata_dir, f"{selected_lang}.traineddata")):
#                 logger.error("Neither Vietnamese nor English language data found")
#                 return ""
        
#         # Use path with forward slashes for consistency
#         tessdata_path = tessdata_dir.replace('\\', '/')
        
#         # Try different PSM modes for better results
#         psm_modes = [6, 4, 3, 1]  # Different page segmentation modes
#         best_text = ""
#         best_length = 0
        
#         for psm in psm_modes:
#             try:
#                 config = f'--tessdata-dir {tessdata_path} -l {selected_lang} --psm {psm} -c preserve_interword_spaces=1'
#                 text = pytesseract.image_to_string(processed_img, config=config)
                
#                 # Choose the result with most content
#                 if len(text.strip()) > best_length:
#                     best_text = text
#                     best_length = len(text.strip())
                    
#             except Exception as e:
#                 logger.warning(f"OCR failed with PSM {psm}: {str(e)}")
#                 continue
        
#         logger.info(f"OCR extracted {best_length} characters from page {page_num} using {selected_lang}")
#         return best_text.strip()
        
#     except Exception as e:
#         logger.error(f"OCR Error for page {page_num}: {str(e)}")
#         return ""

# def process_pdf(file_path: str) -> list:
#     """Xử lý file PDF, trích xuất văn bản bằng OCR."""
#     texts = []
    
#     try:
#         if not os.path.exists(file_path):
#             logger.error(f"File not found: {file_path}")
#             raise FileNotFoundError(f"File not found: {file_path}")
        
#         logger.info(f"Processing PDF: {file_path}")
        
#         # Convert PDF to images
#         try:
#             images = convert_from_path(file_path, dpi=300, fmt='PNG')
#             logger.info(f"Converted PDF to {len(images)} images")
#         except Exception as e:
#             logger.error(f"Failed to convert PDF to images: {str(e)}")
#             return texts
        
#         # Process each page
#         for i, img in enumerate(images):
#             page_num = i + 1
#             logger.info(f"Processing page {page_num}/{len(images)}")
            
#             # Extract text using OCR
#             text = extract_text_with_tesseract(img, page_num)
            
#             if text and text.strip():
#                 # Add page information to the text
                
#                 texts.append(text)
#                 logger.info(f"Successfully extracted text from page {page_num}")
#             else:
#                 logger.warning(f"No text extracted from page {page_num}")
        
#         logger.info(f"Extracted text from {len(texts)} pages of {file_path}")
        
#     except Exception as e:
#         logger.error(f"PDF Processing Error for {file_path}: {str(e)}")
    
#     return texts

# def load_new_documents(file_path: str, metadata) -> list:
#     """Load documents from various file formats"""
#     documents = []
    
#     if not os.path.exists(file_path):
#         logger.error(f"File not found: {file_path}")
#         return documents

#     extension = file_path.lower().split('.')[-1]
#     supported_extensions = {
#         'pdf': 'pdf_ocr',  # Special handling for PDF with OCR
#         'txt': TextLoader,
#         'docx': Docx2txtLoader,
#         'csv': CSVLoader,
#         'xlsx': UnstructuredExcelLoader,
#         'xls': UnstructuredExcelLoader
#     }

#     if extension in supported_extensions:
#         try:
#             logger.info(f"Loading document: {file_path} with extension {extension}")
            
#             if extension == 'pdf':
#                 # Special PDF processing with OCR
#                 texts = process_pdf(file_path)
#                 metadata_dict = metadata.dict(by_alias=True) if hasattr(metadata, 'dict') else metadata
                
#                 for text in texts:
#                     if text and text.strip():
#                         documents.append(Document(page_content=text, metadata=metadata_dict))
                        
#             else:
#                 # Standard document loading
#                 loader = supported_extensions[extension](file_path)
#                 loaded_docs = loader.load()
#                 metadata_dict = metadata.dict(by_alias=True) if hasattr(metadata, 'dict') else metadata
                
#                 for doc in loaded_docs:
#                     documents.append(Document(
#                         page_content=doc.page_content, 
#                         metadata=metadata_dict
#                     ))
            
#             logger.info(f"Loaded {len(documents)} documents from {file_path}")
            
#         except Exception as e:
#             logger.error(f"Error loading {file_path}: {str(e)}")
#     else:
#         logger.warning(f"Unsupported file extension: {extension}")
    
#     return documents


# app/services/document_loader.py
# import os
# import logging
# import cv2
# import numpy as np
# import pytesseract
# import fitz  # PyMuPDF
# from pdf2image import convert_from_path
# from PIL import Image
# from langchain.docstore.document import Document
# from langchain_community.document_loaders import (
#     TextLoader,
#     Docx2txtLoader,
#     CSVLoader,
#     UnstructuredExcelLoader
# )

# logger = logging.getLogger(__name__)

# # Global variables for OCR
# ocr_initialized = False

# def init_tesseract():
#     """Khởi tạo và kiểm tra Tesseract OCR"""
#     global ocr_initialized
    
#     if ocr_initialized:
#         return True
    
#     try:
#         logger.info("Initializing Tesseract OCR...")
        
#         # Kiểm tra tesseract có sẵn không
#         import subprocess
#         try:
#             result = subprocess.run(['tesseract', '--version'], 
#                                   capture_output=True, text=True, timeout=10)
#             if result.returncode == 0:
#                 logger.info(f"Tesseract version: {result.stdout.split()[1]}")
#             else:
#                 logger.error("Tesseract not found in system PATH")
#                 return False
#         except subprocess.TimeoutExpired:
#             logger.error("Tesseract version check timed out")
#             return False
#         except FileNotFoundError:
#             logger.error("Tesseract executable not found")
#             return False
        
#         # Kiểm tra các ngôn ngữ có sẵn
#         try:
#             result = subprocess.run(['tesseract', '--list-langs'], 
#                                   capture_output=True, text=True, timeout=10)
#             if result.returncode == 0:
#                 available_langs = result.stdout.strip().split('\n')[1:]  # Skip header
#                 logger.info(f"Available languages: {available_langs}")
                
#                 # Kiểm tra có tiếng Việt và tiếng Anh không
#                 has_vie = 'vie' in available_langs
#                 has_eng = 'eng' in available_langs
                
#                 if not has_vie and not has_eng:
#                     logger.error("Neither Vietnamese nor English language support found")
#                     return False
                    
#                 logger.info(f"Language support - Vietnamese: {has_vie}, English: {has_eng}")
#             else:
#                 logger.warning("Could not list available languages, proceeding with default")
#         except subprocess.TimeoutExpired:
#             logger.warning("Language list check timed out, proceeding with default")
#         except Exception as e:
#             logger.warning(f"Language check failed: {e}, proceeding with default")
        
#         # Test OCR với hình ảnh đơn giản
#         try:
#             test_img = Image.new('RGB', (200, 50), color='white')
#             from PIL import ImageDraw
#             draw = ImageDraw.Draw(test_img)
#             draw.text((10, 10), "TEST", fill='black')
            
#             # Test với cấu hình cơ bản
#             test_result = pytesseract.image_to_string(test_img, lang='eng')
#             logger.info("OCR test successful")
#         except Exception as e:
#             logger.error(f"OCR test failed: {e}")
#             return False
        
#         ocr_initialized = True
#         logger.info("Tesseract OCR initialized successfully")
#         return True
        
#     except Exception as e:
#         logger.error(f"Failed to initialize Tesseract OCR: {str(e)}")
#         return False

def is_probably_scanned(pdf_path, text_threshold=30, image_area_ratio=0.7):
    """Phân tích PDF để xác định có phải là file scan không"""
    results = []
    doc = None
    
    try:
        if not os.path.exists(pdf_path):
            logger.error(f"PDF file not found: {pdf_path}")
            return results
            
        doc = fitz.open(pdf_path)
        logger.info(f"Analyzing PDF structure: {pdf_path} ({len(doc)} pages)")

        for i, page in enumerate(doc):
            try:
                text = page.get_text().strip()
                image_list = page.get_images(full=True)
                has_text = len(text) >= text_threshold

                # Tính tỷ lệ ảnh trên trang
                img_area_ratio = 0
                page_area = page.rect.width * page.rect.height

                if page_area > 0:  # Tránh chia cho 0
                    for img in image_list:
                        try:
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            if pix:
                                img_width, img_height = pix.width, pix.height
                                img_area_ratio += (img_width * img_height) / page_area
                                pix = None  # Clean up
                        except Exception as e:
                            logger.warning(f"Error processing image on page {i+1}: {str(e)}")
                            continue

                # Nếu text ngắn hoặc ảnh chiếm diện tích lớn → có thể là scan
                is_scanned = (not has_text) or (img_area_ratio > image_area_ratio)
                results.append({
                    "page": i+1,
                    "has_text": has_text,
                    "image_area_ratio": round(img_area_ratio, 2),
                    "is_probably_scan": is_scanned
                })
                
            except Exception as e:
                logger.error(f"Error analyzing page {i+1}: {str(e)}")
                # Add default result for failed page analysis
                results.append({
                    "page": i+1,
                    "has_text": False,
                    "image_area_ratio": 0,
                    "is_probably_scan": True  # Conservative approach
                })

    except Exception as e:
        logger.error(f"Error opening or analyzing PDF {pdf_path}: {str(e)}")
    finally:
        if doc:
            try:
                doc.close()
            except:
                pass
    
    return results

# def extract_text_from_pdf_native(pdf_path):
#     """Trích xuất text trực tiếp từ PDF (không qua OCR)"""
#     texts = []
#     doc = None
    
#     try:
#         if not os.path.exists(pdf_path):
#             logger.error(f"PDF file not found: {pdf_path}")
#             return texts
            
#         doc = fitz.open(pdf_path)
#         logger.info(f"Extracting native text from {len(doc)} pages")
        
#         for page_num in range(len(doc)):
#             try:
#                 page = doc[page_num]
#                 text = page.get_text()
                
#                 if text and text.strip():
#                     texts.append(text.strip())
#                     logger.info(f"Extracted native text from page {page_num + 1}")
                    
#             except Exception as e:
#                 logger.error(f"Error extracting text from page {page_num + 1}: {str(e)}")
#                 continue
            
#         logger.info(f"Successfully extracted native text from {len(texts)} pages")
        
#     except Exception as e:
#         logger.error(f"Error opening PDF for text extraction {pdf_path}: {str(e)}")
#     finally:
#         if doc:
#             try:
#                 doc.close()
#             except:
#                 pass
    
#     return texts

# def preprocess_image_for_ocr(image):
#     """Tiền xử lý hình ảnh để cải thiện OCR"""
#     try:
#         # Convert PIL Image to numpy array
#         img_array = np.array(image)
        
#         # Convert RGB to BGR for OpenCV
#         if len(img_array.shape) == 3:
#             img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
#             gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
#         else:
#             gray = img_array
        
#         # Apply CLAHE (Contrast Limited Adaptive Histogram Equalization)
#         clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
#         enhanced = clahe.apply(gray)
        
#         # Denoise
#         denoised = cv2.medianBlur(enhanced, 3)
        
#         # Threshold
#         _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
#         return Image.fromarray(thresh)
        
#     except Exception as e:
#         logger.warning(f"Image preprocessing failed: {str(e)}, using original image")
#         return image

# def extract_text_with_tesseract(image, page_num=1):
#     """Trích xuất văn bản từ hình ảnh bằng Tesseract OCR"""
#     if not init_tesseract():
#         logger.error("Cannot initialize Tesseract OCR")
#         return ""
    
#     try:
#         # Preprocess image
#         processed_img = preprocess_image_for_ocr(image)
        
#         # Kiểm tra ngôn ngữ có sẵn
#         import subprocess
#         available_langs = []
#         try:
#             result = subprocess.run(['tesseract', '--list-langs'], 
#                                   capture_output=True, text=True, timeout=5)
#             if result.returncode == 0:
#                 available_langs = result.stdout.strip().split('\n')[1:]
#         except:
#             available_langs = ['eng']  # fallback
        
#         # Chọn ngôn ngữ ưu tiên
#         if 'vie' in available_langs:
#             selected_lang = 'vie+eng'  # Combine Vietnamese and English
#         elif 'eng' in available_langs:
#             selected_lang = 'eng'
#         else:
#             selected_lang = 'eng'  # Default fallback
        
#         # Try different PSM modes for better results
#         psm_modes = [6, 4, 3, 1]  # Different page segmentation modes
#         best_text = ""
#         best_length = 0
        
#         for psm in psm_modes:
#             try:
#                 config = f'--psm {psm} -c preserve_interword_spaces=1'
#                 text = pytesseract.image_to_string(processed_img, lang=selected_lang, config=config)
                
#                 # Choose the result with most content
#                 if len(text.strip()) > best_length:
#                     best_text = text
#                     best_length = len(text.strip())
                    
#             except Exception as e:
#                 logger.warning(f"OCR failed with PSM {psm}: {str(e)}")
#                 continue
        
#         logger.info(f"OCR extracted {best_length} characters from page {page_num} using {selected_lang}")
#         return best_text.strip()
        
#     except Exception as e:
#         logger.error(f"OCR Error for page {page_num}: {str(e)}")
#         return ""

def process_pdf_smart(file_path: str) -> list:
    """
    Xử lý file PDF thông minh: 
    - Phân tích cấu trúc để xác định cần OCR hay không
    - Sử dụng text extraction trực tiếp cho PDF thông thường
    - Chỉ sử dụng OCR cho PDF scan
    """
    texts = []
    doc = None
    
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Processing PDF: {file_path}")
        
        # Phân tích cấu trúc PDF
        scan_analysis = is_probably_scanned(file_path)
        
        if not scan_analysis:
            logger.error("Failed to analyze PDF structure")
            return texts
        
        # Phân loại trang: scan vs native text
        scan_pages = []
        native_pages = []
        
        for page_info in scan_analysis:
            if page_info["is_probably_scan"]:
                scan_pages.append(page_info["page"])
            else:
                native_pages.append(page_info["page"])
        
        logger.info(f"PDF Analysis: {len(native_pages)} native text pages, {len(scan_pages)} scan pages")
        
        # Xử lý native text pages
        if native_pages:
            try:
                logger.info("Extracting native text from PDF...")
                doc = fitz.open(file_path)
                
                for page_num in native_pages:
                    try:
                        page_index = page_num - 1  # Convert to 0-based index
                        if page_index < len(doc):
                            page = doc[page_index]
                            text = page.get_text().strip()
                            if text:
                                texts.append(text)
                                logger.info(f"Native text extracted from page {page_num}")
                    except Exception as e:
                        logger.warning(f"Error extracting native text from page {page_num}: {str(e)}")
                        continue
                        
            except Exception as e:
                logger.error(f"Error processing native text pages: {str(e)}")
            finally:
                if doc:
                    try:
                        doc.close()
                        doc = None
                    except:
                        pass
        
        # Xử lý scan pages với OCR
        if scan_pages:
            logger.info(f"Processing {len(scan_pages)} scan pages with OCR...")
            
            try:
                # Convert PDF to images (tất cả trang)
                images = convert_from_path(file_path, dpi=300, fmt='PNG')
                logger.info(f"Converted PDF to {len(images)} images")
                
                # Process only scan pages
                for page_num in scan_pages:
                    try:
                        img_index = page_num - 1  # Convert to 0-based index
                        if img_index < len(images):
                            img = images[img_index]
                            logger.info(f"Processing scan page {page_num} with OCR")
                            
                            # Extract text using OCR
                            text = extract_text_with_tesseract(img, page_num)
                            
                            if text and text.strip():
                                texts.append(text)
                                logger.info(f"OCR text extracted from page {page_num}")
                            else:
                                logger.warning(f"No text extracted from scan page {page_num}")
                        else:
                            logger.warning(f"Image index {img_index} out of range for page {page_num}")
                        
                    except Exception as e:
                        logger.error(f"Error processing scan page {page_num}: {str(e)}")
                        continue
                        
            except Exception as e:
                logger.error(f"Failed to convert PDF to images: {str(e)}")
                # Don't return here, we might have extracted some native text
        
        logger.info(f"Successfully processed PDF: {len(texts)} pages with text extracted")
        
    except FileNotFoundError:
        raise  # Re-raise file not found error
    except Exception as e:
        logger.error(f"PDF Processing Error for {file_path}: {str(e)}")
    finally:
        if doc:
            try:
                doc.close()
            except:
                pass
    
    return texts

# def load_new_documents(file_path: str, metadata) -> list:
#     """Load documents from various file formats with smart PDF processing"""
#     documents = []
    
#     try:
#         if not os.path.exists(file_path):
#             logger.error(f"File not found: {file_path}")
#             return documents

#         extension = file_path.lower().split('.')[-1]
#         supported_extensions = {
#             'pdf': 'pdf_smart',  # Smart PDF processing
#             'txt': TextLoader,
#             'docx': Docx2txtLoader,
#             'csv': CSVLoader,
#             'xlsx': UnstructuredExcelLoader,
#             'xls': UnstructuredExcelLoader
#         }

#         if extension in supported_extensions:
#             try:
#                 logger.info(f"Loading document: {file_path} with extension {extension}")
                
#                 if extension == 'pdf':
#                     try:
#                         # Smart PDF processing
#                         texts = process_pdf_smart(file_path)
#                         metadata_dict = metadata.dict(by_alias=True) if hasattr(metadata, 'dict') else metadata
                        
#                         for text in texts:
#                             try:
#                                 if text and text.strip():
#                                     documents.append(Document(page_content=text, metadata=metadata_dict))
#                             except Exception as e:
#                                 logger.warning(f"Error creating document from text: {str(e)}")
#                                 continue
                                
#                     except Exception as e:
#                         logger.error(f"Error processing PDF {file_path}: {str(e)}")
                        
#                 else:
#                     try:
#                         # Standard document loading
#                         loader = supported_extensions[extension](file_path)
#                         loaded_docs = loader.load()
#                         metadata_dict = metadata.dict(by_alias=True) if hasattr(metadata, 'dict') else metadata
                        
#                         for doc in loaded_docs:
#                             try:
#                                 documents.append(Document(
#                                     page_content=doc.page_content, 
#                                     metadata=metadata_dict
#                                 ))
#                             except Exception as e:
#                                 logger.warning(f"Error creating document from loaded doc: {str(e)}")
#                                 continue
                                
#                     except Exception as e:
#                         logger.error(f"Error loading document with standard loader: {str(e)}")
                
#                 logger.info(f"Successfully loaded {len(documents)} documents from {file_path}")
                
#             except Exception as e:
#                 logger.error(f"Error processing file {file_path}: {str(e)}")
#         else:
#             logger.warning(f"Unsupported file extension: {extension} for file {file_path}")
            
#     except Exception as e:
#         logger.error(f"Unexpected error in load_new_documents for {file_path}: {str(e)}")
    
#     return documents

import os
import logging
import cv2
import numpy as np
import pytesseract
import fitz  # PyMuPDF
import pandas as pd
from pdf2image import convert_from_path
from PIL import Image
from langchain.docstore.document import Document
from langchain_community.document_loaders import (
    TextLoader,
    Docx2txtLoader,
    CSVLoader,
    UnstructuredExcelLoader
)
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Global variables for OCR
ocr_initialized = False

def init_tesseract():
    """Khởi tạo và kiểm tra Tesseract OCR"""
    global ocr_initialized
    
    if ocr_initialized:
        return True
    
    try:
        logger.info("Initializing Tesseract OCR...")
        
        # Kiểm tra tesseract có sẵn không
        import subprocess
        try:
            result = subprocess.run(['tesseract', '--version'], 
                                  capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                logger.info(f"Tesseract version: {result.stdout.split()[1]}")
            else:
                logger.error("Tesseract not found in system PATH")
                return False
        except subprocess.TimeoutExpired:
            logger.error("Tesseract version check timed out")
            return False
        except FileNotFoundError:
            logger.error("Tesseract executable not found")
            return False
        
        # Kiểm tra các ngôn ngữ có sẵn
        try:
            result = subprocess.run(['tesseract', '--list-langs'], 
                                  capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                available_langs = result.stdout.strip().split('\n')[1:]  # Skip header
                
                # Kiểm tra có tiếng Việt và tiếng Anh không
                has_vie = 'vie' in available_langs
                has_eng = 'eng' in available_langs
                
                if not has_vie and not has_eng:
                    return False
                    
            else:
                logger.warning("Could not list available languages, proceeding with default")
        except subprocess.TimeoutExpired:
            logger.warning("Language list check timed out, proceeding with default")
        except Exception as e:
            logger.warning(f"Language check failed: {e}, proceeding with default")
        
        # Test OCR với hình ảnh đơn giản
        try:
            test_img = Image.new('RGB', (200, 50), color='white')
            from PIL import ImageDraw
            draw = ImageDraw.Draw(test_img)
            draw.text((10, 10), "TEST", fill='black')
            
            # Test với cấu hình cơ bản
            test_result = pytesseract.image_to_string(test_img, lang='eng')
        except Exception as e:
            return False
        
        ocr_initialized = True
        return True
        
    except Exception as e:
        return False

def is_probably_scanned_first_page(pdf_path, text_threshold=30, image_area_ratio=0.7):
    """Phân tích trang đầu tiên của PDF để xác định có phải là file scan không"""
    result = {"page": 1, "has_text": False, "image_area_ratio": 0, "is_probably_scan": True}
    doc = None
    
    try:
        if not os.path.exists(pdf_path):
            return result
            
        doc = fitz.open(pdf_path)
        if len(doc) == 0:
            return result
            
        page = doc[0]  # Chỉ kiểm tra trang đầu tiên
        
        try:
            text = page.get_text().strip()
            image_list = page.get_images(full=True)
            has_text = len(text) >= text_threshold

            # Tính tỷ lệ ảnh trên trang
            img_area_ratio = 0
            page_area = page.rect.width * page.rect.height

            if page_area > 0:  # Tránh chia cho 0
                for img in image_list:
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        if pix:
                            img_width, img_height = pix.width, pix.height
                            img_area_ratio += (img_width * img_height) / page_area
                            pix = None  # Clean up
                    except Exception as e:
                        logger.warning(f"Error processing image on page 1: {str(e)}")
                        continue

            # Nếu text ngắn hoặc ảnh chiếm diện tích lớn → có thể là scan
            is_scanned = (not has_text) or (img_area_ratio > image_area_ratio)
            result = {
                "page": 1,
                "has_text": has_text,
                "image_area_ratio": round(img_area_ratio, 2),
                "is_probably_scan": is_scanned
            }
                
        except Exception as e:
            logger.error(f"Error analyzing page 1: {str(e)}")
            # Mặc định coi là scan nếu phân tích thất bại

    except Exception as e:
        logger.error(f"Error opening or analyzing PDF {pdf_path}: {str(e)}")
    finally:
        if doc:
            try:
                doc.close()
            except:
                pass
    
    return result

def extract_text_from_pdf_native(pdf_path):
    """Trích xuất text trực tiếp từ PDF (không qua OCR)"""
    texts = []
    doc = None
    
    try:
        if not os.path.exists(pdf_path):
            return texts
            
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            try:
                page = doc[page_num]
                text = page.get_text()
                
                if text and text.strip():
                    texts.append(text.strip())
                    
            except Exception as e:
                logger.error(f"Error extracting text from page {page_num + 1}: {str(e)}")
                continue
            
        logger.info(f"Successfully extracted native text from {len(texts)} pages")
        
    except Exception as e:
        logger.error(f"Error opening PDF for text extraction {pdf_path}: {str(e)}")
    finally:
        if doc:
            try:
                doc.close()
            except:
                pass
    
    return texts

def preprocess_image_for_ocr(image):
    """Tiền xử lý hình ảnh để cải thiện OCR"""
    try:
        # Convert PIL Image to numpy array
        img_array = np.array(image)
        
        # Convert RGB to BGR for OpenCV
        if len(img_array.shape) == 3:
            img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        else:
            gray = img_array
        
        # Apply CLAHE (Contrast Limited Adaptive Histogram Equalization)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        
        # Denoise
        denoised = cv2.medianBlur(enhanced, 3)
        
        # Threshold
        _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return Image.fromarray(thresh)
        
    except Exception as e:
        logger.warning(f"Image preprocessing failed: {str(e)}, using original image")
        return image

def extract_text_with_tesseract(image, page_num=1):
    """Trích xuất văn bản từ hình ảnh bằng Tesseract OCR"""
    if not init_tesseract():
        logger.error("Cannot initialize Tesseract OCR")
        return ""
    
    try:
        # Preprocess image
        processed_img = preprocess_image_for_ocr(image)
        
        # Kiểm tra ngôn ngữ có sẵn
        import subprocess
        available_langs = []
        try:
            result = subprocess.run(['tesseract', '--list-langs'], 
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                available_langs = result.stdout.strip().split('\n')[1:]
        except:
            available_langs = ['eng']  # fallback
        
        # Chọn ngôn ngữ ưu tiên
        if 'vie' in available_langs:
            selected_lang = 'vie+eng'  # Combine Vietnamese and English
        elif 'eng' in available_langs:
            selected_lang = 'eng'
        else:
            selected_lang = 'eng'  # Default fallback
        
        # Try different PSM modes for better results
        psm_modes = [6, 4, 3, 1]  # Different page segmentation modes
        best_text = ""
        best_length = 0
        
        for psm in psm_modes:
            try:
                config = f'--psm {psm} -c preserve_interword_spaces=1'
                text = pytesseract.image_to_string(processed_img, lang=selected_lang, config=config)
                
                # Choose the result with most content
                if len(text.strip()) > best_length:
                    best_text = text
                    best_length = len(text.strip())
                    
            except Exception as e:
                logger.warning(f"OCR failed with PSM {psm}: {str(e)}")
                continue
        
        logger.info(f"OCR extracted {best_length} characters from page {page_num} using {selected_lang}")
        return best_text.strip()
        
    except Exception as e:
        logger.error(f"OCR Error for page {page_num}: {str(e)}")
        return ""

# def process_pdf_smart(file_path: str) -> list:
#     """
#     Xử lý file PDF thông minh:
#     - Kiểm tra trang đầu tiên để quyết định phương pháp xử lý (text hoặc OCR) cho toàn bộ file
#     """
#     texts = []
#     doc = None
    
#     try:
#         if not os.path.exists(file_path):
#             logger.error(f"File not found: {file_path}")
#             raise FileNotFoundError(f"File not found: {file_path}")
        
#         logger.info(f"Processing PDF: {file_path}")
        
#         # Phân tích trang đầu tiên
#         first_page_analysis = is_probably_scanned_first_page(file_path)
        
#         if not first_page_analysis:
#             logger.error("Failed to analyze first page")
#             return texts
        
#         is_scanned = first_page_analysis["is_probably_scan"]
#         logger.info(f"First page analysis: {'scanned' if is_scanned else 'native text'}")
        
#         if not is_scanned:
#             # Trang đầu tiên là text -> trích xuất văn bản gốc cho toàn bộ file
#             logger.info("Processing entire PDF as native text...")
#             texts = extract_text_from_pdf_native(file_path)
#         else:
#             # Trang đầu tiên là scan -> áp dụng OCR cho toàn bộ file
#             logger.info("Processing entire PDF with OCR...")
#             try:
#                 images = convert_from_path(file_path, dpi=300, fmt='PNG')
#                 logger.info(f"Converted PDF to {len(images)} images")
                
#                 for page_num, img in enumerate(images, 1):
#                     try:
#                         logger.info(f"Processing page {page_num} with OCR")
#                         text = extract_text_with_tesseract(img, page_num)
#                         if text and text.strip():
#                             texts.append(text)
#                             logger.info(f"OCR text extracted from page {page_num}")
#                         else:
#                             logger.warning(f"No text extracted from page {page_num}")
#                     except Exception as e:
#                         logger.error(f"Error processing page {page_num}: {str(e)}")
#                         continue
#             except Exception as e:
#                 logger.error(f"Failed to convert PDF to images: {str(e)}")
        
#         logger.info(f"Successfully processed PDF: {len(texts)} pages with text extracted")
        
#     except FileNotFoundError:
#         raise
#     except Exception as e:
#         logger.error(f"PDF Processing Error for {file_path}: {str(e)}")
#     finally:
#         if doc:
#             try:
#                 doc.close()
#             except:
#                 pass
    
#     return texts

def extract_docx_with_tables(docx_path, table_sep="\t"):
    """
    Trích xuất văn bản và bảng từ file DOCX
    """
    try:
        doc = DocxDocument(docx_path)
        full_text = []
        
        # Lấy tất cả paragraphs và tables theo thứ tự xuất hiện
        for element in doc.element.body:
            if element.tag.endswith('}p'):  # paragraph
                # Tìm paragraph tương ứng
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            full_text.append(text)
                        break
            
            elif element.tag.endswith('}tbl'):  # table
                # Tìm table tương ứng
                for table in doc.tables:
                    if table._element == element:
                        # full_text.append("\n--- TABLE ---")
                        
                        # Trích xuất bảng
                        table_text = []
                        for row in table.rows:
                            row_text = table_sep.join(cell.text.strip() for cell in row.cells)
                            if row_text.strip():  # Chỉ thêm dòng không rỗng
                                table_text.append(row_text)
                        
                        if table_text:
                            full_text.extend(table_text)
                        
                        # full_text.append("--- END TABLE ---\n")
                        break
        
        # Kết hợp tất cả text
        combined_text = "\n".join(full_text)
        logger.info(f"Extracted {len(combined_text)} characters from DOCX with tables")
        
        return combined_text
        
    except Exception as e:
        logger.error(f"Error extracting DOCX with tables: {str(e)}")
        # Fallback: chỉ lấy text thông thường
        try:
            doc = DocxDocument(docx_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            logger.info(f"Fallback: Extracted {len(text)} characters from DOCX (text only)")
            return text
        except Exception as fallback_e:
            logger.error(f"Fallback extraction also failed: {str(fallback_e)}")
            return ""

def split_content_into_chunks(content: str, chunk_size: int = 6000, chunk_overlap: int = 600) -> list:
    """
    Chia nội dung thành các chunk nhỏ hơn để tránh quá dài
    
    Args:
        content: Nội dung cần chia
        chunk_size: Kích thước tối đa của mỗi chunk (ký tự)
        chunk_overlap: Số ký tự overlap giữa các chunk
    
    Returns:
        List các chunk đã được chia
    """
    if not content or not content.strip():
        return []
    
    content = content.strip()
    
    # Nếu nội dung ngắn hơn chunk_size, trả về nguyên văn
    if len(content) <= chunk_size:
        return [content]
    
    chunks = []
    start = 0
    
    while start < len(content):
        # Tính toán end position
        end = start + chunk_size
        
        if end >= len(content):
            # Chunk cuối cùng
            chunks.append(content[start:])
            break
        
        # Tìm điểm cắt tự nhiên (ưu tiên paragraph, sentence)
        chunk_end = end
        
        # Tìm xuống dòng gần nhất (paragraph break)
        for i in range(end, max(start + chunk_size // 2, start + 100), -1):
            if i < len(content) and content[i] == '\n' and content[i-1] == '\n':
                chunk_end = i
                break
        else:
            # Nếu không tìm thấy paragraph break, tìm xuống dòng đơn
            for i in range(end, max(start + chunk_size // 2, start + 100), -1):
                if i < len(content) and content[i] == '\n':
                    chunk_end = i
                    break
            else:
                # Nếu không tìm thấy xuống dòng, tìm dấu câu
                for i in range(end, max(start + chunk_size // 2, start + 100), -1):
                    if i < len(content) and content[i] in '.!?':
                        chunk_end = i + 1
                        break
                else:
                    # Nếu không tìm thấy dấu câu, tìm khoảng trắng
                    for i in range(end, max(start + chunk_size // 2, start + 100), -1):
                        if i < len(content) and content[i] == ' ':
                            chunk_end = i
                            break
                    else:
                        # Cuối cùng, cắt cứng
                        chunk_end = end
        
        # Thêm chunk
        chunk = content[start:chunk_end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Cập nhật start position với overlap
        start = max(chunk_end - chunk_overlap, start + 1)
    
    return chunks

def load_excel_or_csv(file_path: str) -> str:
    """
    Đọc file Excel (.xlsx, .xls) hoặc CSV và trả về nội dung dạng chuỗi.
    
    Args:
        file_path: Đường dẫn đến file Excel hoặc CSV.
    
    Returns:
        Chuỗi nội dung được trích xuất từ file.
    """
    try:
        extension = file_path.lower().split('.')[-1]
        
        if extension in ['xlsx', 'xls']:
            # Đọc file Excel bằng pandas
            df = pd.read_excel(file_path, engine='openpyxl')
        elif extension == 'csv':
            # Đọc file CSV bằng pandas
            df = pd.read_csv(file_path, encoding='utf-8')
        else:
            logger.error(f"Định dạng file không được hỗ trợ: {extension}")
            return ""
        
        # Chuyển nội dung thành chuỗi, giữ cấu trúc dòng
        rows_content = []
        for index, row in df.iterrows():
            row_content = [str(cell) for cell in row if pd.notna(cell)]  # Bỏ qua giá trị NaN
            formatted_row = " | ".join(row_content)  # Kết hợp các cột bằng dấu phân cách
            rows_content.append(formatted_row)
        
        # Kết hợp tất cả dòng thành một chuỗi, thêm khoảng cách giữa các dòng
        combined_content = "\n\n".join(rows_content)
        logger.info(f"Đã trích xuất {len(rows_content)} dòng từ file {file_path}")
        return combined_content

    except Exception as e:
        logger.error(f"Lỗi khi đọc file {file_path}: {str(e)}")
        return ""


def load_new_documents(file_path: str, metadata, chunk_size: int =6000, chunk_overlap: int = 600) -> list:
    """
    Load documents from various file formats with content splitting
    
    Args:
        file_path: Đường dẫn đến file
        metadata: Metadata cho document
        chunk_size: Kích thước tối đa của mỗi chunk (ký tự)
        chunk_overlap: Số ký tự overlap giữa các chunk
    
    Returns:
        List các Document đã được chia nhỏ
    """
    documents = []
    
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return documents

        extension = file_path.lower().split('.')[-1]
        supported_extensions = {
            'pdf': 'pdf_smart',  # Smart PDF processing
            'txt': TextLoader,
            'docx': 'docx_with_tables',  # Enhanced DOCX processing
            'csv': 'pandas_loader',
            'xlsx': 'pandas_loader',
            'xls': 'pandas_loader'
        }

        if extension in supported_extensions:
            try:
                logger.info(f"Loading document: {file_path} with extension {extension}")
                combined_content = ""  # Biến để kết hợp nội dung
                
                if extension == 'pdf':
                    try:
                        # Smart PDF processing - kết hợp tất cả text
                        texts = process_pdf_smart(file_path)
                        
                        for text in texts:
                            if text and text.strip():
                                combined_content += text + "\n\n"  # Thêm khoảng cách giữa các trang
                                
                    except Exception as e:
                        logger.error(f"Error processing PDF {file_path}: {str(e)}")
                
                elif extension == 'docx':
                    try:
                        # Enhanced DOCX processing with table support
                        combined_content = extract_docx_with_tables(file_path)
                            
                    except Exception as e:
                        logger.error(f"Error processing DOCX {file_path}: {str(e)}")
                        
                else:
                    try:
                        # Xử lý file CSV và Excel bằng pandas
                        if extension in ['csv', 'xlsx', 'xls']:
                            combined_content = load_excel_or_csv(file_path)
                        else:
                            # Xử lý file txt bằng TextLoader
                            loader = supported_extensions[extension](file_path)
                            loaded_docs = loader.load()
                            for doc in loaded_docs:
                                if doc.page_content and doc.page_content.strip():
                                    combined_content += doc.page_content + "\n\n"
                                
                    except Exception as e:
                        logger.error(f"Error loading document with standard loader: {str(e)}")
                
                # Chia nội dung thành các chunk nhỏ hơn
                if combined_content and combined_content.strip():
                    chunks = split_content_into_chunks(combined_content.strip(), chunk_size, chunk_overlap)
                    
                    # Tạo Document cho mỗi chunk
                    logger.info(f"Splitting content into {len(chunks)} chunks")
                    for i, chunk in enumerate(chunks):
                        if chunk and chunk.strip():
                            # Tạo metadata cho từng chunk
                            chunk_metadata = metadata.copy() if isinstance(metadata, dict) else metadata.dict(by_alias=True) if hasattr(metadata, 'dict') else metadata
                            
                            documents.append(Document(
                                page_content=chunk.strip(),
                                metadata=chunk_metadata
                            ))
                    
                    logger.info(f"Successfully created {len(documents)} document chunks from {file_path}")
                else:
                    logger.warning(f"No content extracted from file: {file_path}")
                
            except Exception as e:
                logger.error(f"Error processing file {file_path}: {str(e)}")
        else:
            logger.warning(f"Unsupported file extension: {extension} for file {file_path}")
            
    except Exception as e:
        logger.error(f"Unexpected error in load_new_documents for {file_path}: {str(e)}")
    
    return documents



def load_new_documents_simple(file_path: str, metadata, chunk_size: int = 2000) -> list:
    """
    Phiên bản đơn giản hơn - chỉ chia theo kích thước cố định, không có metadata bổ sung
    
    Args:
        file_path: Đường dẫn đến file  
        metadata: Metadata gốc
        chunk_size: Kích thước tối đa của mỗi chunk
    
    Returns:
        List các Document đã được chia nhỏ (không có metadata bổ sung)
    """
    documents = []
    
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return documents

        extension = file_path.lower().split('.')[-1]
        supported_extensions = {
            'pdf': 'pdf_smart',
            'txt': TextLoader,
            'docx': 'docx_with_tables',
            'csv': CSVLoader,
            'xlsx': UnstructuredExcelLoader,
            'xls': UnstructuredExcelLoader
        }

        if extension in supported_extensions:
            try:
                logger.info(f"Loading document: {file_path} with extension {extension}")
                combined_content = ""
                
                if extension == 'pdf':
                    try:
                        texts = process_pdf_smart(file_path)
                        for text in texts:
                            if text and text.strip():
                                combined_content += text + "\n\n"
                    except Exception as e:
                        logger.error(f"Error processing PDF {file_path}: {str(e)}")
                
                elif extension == 'docx':
                    try:
                        combined_content = extract_docx_with_tables(file_path)
                    except Exception as e:
                        logger.error(f"Error processing DOCX {file_path}: {str(e)}")
                        
                else:
                    try:
                        loader = supported_extensions[extension](file_path)
                        loaded_docs = loader.load()
                        for doc in loaded_docs:
                            if doc.page_content and doc.page_content.strip():
                                combined_content += doc.page_content + "\n\n"
                    except Exception as e:
                        logger.error(f"Error loading document with standard loader: {str(e)}")
                
                # Chia nội dung thành chunks đơn giản (cứ mỗi chunk_size ký tự)
                if combined_content and combined_content.strip():
                    content = combined_content.strip()
                    
                    # Sử dụng metadata gốc không thay đổi
                    original_metadata = metadata.dict(by_alias=True) if hasattr(metadata, 'dict') else metadata
                    
                    # Chia thành chunks
                    for i in range(0, len(content), chunk_size):
                        chunk = content[i:i + chunk_size]
                        if chunk and chunk.strip():
                            documents.append(Document(
                                page_content=chunk.strip(),
                                metadata=original_metadata  # Metadata gốc, không thêm gì
                            ))
                    
                    logger.info(f"Successfully created {len(documents)} document chunks from {file_path}")
                else:
                    logger.warning(f"No content extracted from file: {file_path}")
                
            except Exception as e:
                logger.error(f"Error processing file {file_path}: {str(e)}")
        else:
            logger.warning(f"Unsupported file extension: {extension} for file {file_path}")
            
    except Exception as e:
        logger.error(f"Unexpected error in load_new_documents for {file_path}: {str(e)}")
    
    return documents